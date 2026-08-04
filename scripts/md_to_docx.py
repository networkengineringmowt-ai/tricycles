import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from bs4 import BeautifulSoup

def convert_md_to_docx(md_path, docx_path):
    print(f"Converting {md_path} to {docx_path}...")
    
    if not os.path.exists(md_path):
        print("Markdown file not found.")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # Convert markdown to html, then parse with BeautifulSoup
    html = markdown.markdown(md_text, extensions=['tables'])
    soup = BeautifulSoup(html, 'html.parser')

    doc = Document()
    
    # Configure styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for element in soup:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'h4':
            doc.add_heading(element.text, level=4)
        elif element.name == 'p':
            # Check if it contains an image
            img = element.find('img')
            if img:
                try:
                    img_path = img['src']
                    if img_path.startswith('../../'):
                        img_path = img_path.replace('../../', '')
                    # Add image
                    if os.path.exists(img_path):
                        doc.add_picture(img_path, width=Inches(6.0))
                        last_par = doc.paragraphs[-1]
                        last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        doc.add_paragraph(f"[Image Missing: {img_path}]")
                except Exception as e:
                    doc.add_paragraph(f"[Error loading image: {str(e)}]")
            else:
                doc.add_paragraph(element.text)
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.text, style='List Number')
        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=max(len(row.find_all(['th', 'td'])) for row in rows))
                table.style = 'Table Grid'
                for i, row in enumerate(rows):
                    cols = row.find_all(['th', 'td'])
                    for j, col in enumerate(cols):
                        if j < len(table.columns):
                            table.cell(i, j).text = col.text

    doc.save(docx_path)
    print(f"Successfully generated Word Document: {docx_path} (Size: {os.path.getsize(docx_path)/1024:.2f} KB)")

if __name__ == "__main__":
    convert_md_to_docx("Final_Thesis.md", "Final_Thesis.docx")
